"""
PAC Document Generator Utility

This module provides functionality to generate PAC (Preliminary Acceptance Certificate)
documents by modifying a Word template with site-specific information.
"""

import os
import re
import zipfile
from io import BytesIO
from typing import Optional
from docx import Document


def extract_site_numbers_from_link(link_id: str) -> str:
    """
    Extract site numbers from Link ID.

    Example: "ABBS2161-ABBS0162" -> "2161-0162"

    Args:
        link_id: The link ID string (e.g., "ABBS2161-ABBS0162")

    Returns:
        Combined site numbers with hyphen
    """
    # Extract all groups of consecutive digits
    numbers = re.findall(r'\d+', link_id)
    if len(numbers) >= 2:
        # Take the last two groups (site numbers)
        return f"{numbers[-2]}-{numbers[-1]}"
    elif len(numbers) == 1:
        # If only one number, use it
        return numbers[0]
    return "0000"


def modify_pac_template(
    template_path: str,
    site_id: str,
    project_name: str,
    project_po: str,
    link_id: str,
    po_line_number: str = "1",
    model_name: str = "Implementation services - New Site",
    output_path: Optional[str] = None
) -> BytesIO:
    """
    Modify PAC template with site-specific information.

    Args:
        template_path: Path to the Word template file (.docx format required)
        site_id: Site ID (e.g., "10_74_124_41")
        project_name: Project name
        project_po: Project PO number
        link_id: Link ID (e.g., "ABBS2161-ABBS0162") for certificate number
        po_line_number: PO line number(s) - default "1" for MW, can be multiple for RAN
        model_name: Model name from Implementation services row
        output_path: Optional path to save the modified document

    Returns:
        BytesIO object containing the modified Word document

    Raises:
        ValueError: If template is not in .docx format
    """
    # Check if template is .docx format
    if not template_path.endswith('.docx'):
        raise ValueError(
            f"Template must be in .docx format. Found: {template_path}\n"
            "Please convert the .doc file to .docx format using Microsoft Word:\n"
            "1. Open the file in Word\n"
            "2. File → Save As\n"
            "3. Choose 'Word Document (*.docx)' format"
        )

    # Extract site number for certificate from link ID
    site_number = extract_site_numbers_from_link(link_id)
    certificate_number = f"#{site_number}/R4"

    # Project description
    project_description = f"Zain / {project_name}, TI Service"

    # Load the template
    try:
        doc = Document(template_path)
        print(f"[PAC_GEN] Successfully loaded template: {template_path}")
    except Exception as e:
        raise ValueError(f"Failed to load template file: {str(e)}")

    # Replacement mappings - match the actual text in the template with spaces
    replacements = {
        "Preliminary Acceptance Certificate # 2876 / R4": f"Preliminary Acceptance Certificate {certificate_number}",
        "# 2876 / R4": certificate_number,
        "JED2876": site_id,
        "ZAIN / SOPHIA 4 , TI Service": project_description,
    }

    print(f"[PAC_GEN] Site ID: {site_id}")
    print(f"[PAC_GEN] Certificate Number: {certificate_number}")
    print(f"[PAC_GEN] Project Description: {project_description}")
    print(f"[PAC_GEN] Project PO: {project_po}")
    print(f"[PAC_GEN] PO Line Number: {po_line_number}")
    print(f"[PAC_GEN] Model Name: {model_name}")
    print(f"[PAC_GEN] Replacements to make: {replacements}")

    # Track what we find
    replacements_made = {}
    for key in replacements.keys():
        replacements_made[key] = 0

    # Helper function to replace text across runs (handles text split across runs)
    def replace_in_paragraph(paragraph, old_text, new_text):
        """Replace text in paragraph, handling text that may be split across runs."""
        if old_text not in paragraph.text:
            return False

        # Try simple replacement first (text in single run)
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return True

        # Text is split across runs - use a smarter approach
        # Build a map of character positions to runs
        full_text = paragraph.text
        if old_text not in full_text:
            return False

        print(f"[PAC_GEN] Text '{old_text}' is split across {len(paragraph.runs)} runs")

        # Find the position of old_text
        start_pos = full_text.find(old_text)
        end_pos = start_pos + len(old_text)

        # Find which runs contain the old text
        current_pos = 0
        affected_runs = []
        for i, run in enumerate(paragraph.runs):
            run_start = current_pos
            run_end = current_pos + len(run.text)

            # Check if this run overlaps with our target text
            if run_start < end_pos and run_end > start_pos:
                overlap_start = max(0, start_pos - run_start)
                overlap_end = min(len(run.text), end_pos - run_start)
                affected_runs.append((i, run, overlap_start, overlap_end, run_start, run_end))

            current_pos = run_end

        if not affected_runs:
            return False

        # Replace text in affected runs
        if len(affected_runs) == 1:
            # Simple case - text is in one run
            i, run, overlap_start, overlap_end, _, _ = affected_runs[0]
            run.text = run.text[:overlap_start] + new_text + run.text[overlap_end:]
        else:
            # Complex case - text spans multiple runs
            # Replace in first run and clear others
            first_i, first_run, first_start, first_end, first_run_start, _ = affected_runs[0]
            prefix = first_run.text[:first_start]

            # Get suffix from last run
            last_i, last_run, last_start, last_end, _, last_run_end = affected_runs[-1]
            suffix = last_run.text[last_end:]

            # Update first run with prefix + new_text + suffix
            first_run.text = prefix + new_text + suffix

            # Clear middle and last runs that were part of the old text
            for i, run, _, _, _, _ in affected_runs[1:]:
                run.text = ""

        return True

    # Function to replace text in paragraphs
    def replace_text_in_paragraph(paragraph, replacements):
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                print(f"[PAC_GEN] Found '{old_text}' in paragraph: {paragraph.text[:100]}")
                if replace_in_paragraph(paragraph, old_text, new_text):
                    replacements_made[old_text] += 1
                    print(f"[PAC_GEN] Replaced '{old_text}' with '{new_text}'")

    # Special handler for PO line number (handles varying whitespace)
    def replace_po_line_number(paragraph, new_po_line):
        """Replace PO line number with flexible whitespace matching."""
        full_text = paragraph.text
        # Pattern matches "PO line number" followed by any whitespace, ":", more whitespace, and a number
        po_pattern = re.search(r'(PO line number\s*:\s*)(\d+)', full_text, re.IGNORECASE)
        if po_pattern:
            old_value = po_pattern.group(2)
            old_full = po_pattern.group(0)
            new_full = po_pattern.group(1) + new_po_line

            print(f"[PAC_GEN] Found PO line number: '{old_full}' in paragraph: {full_text[:100]}")
            if replace_in_paragraph(paragraph, old_full, new_full):
                print(f"[PAC_GEN] Replaced PO line '{old_value}' with '{new_po_line}'")
                return True
        return False

    # Special handler for Model Name (handles varying whitespace)
    def replace_model_name(paragraph, new_model_name):
        """Replace Model Name with flexible whitespace matching."""
        full_text = paragraph.text
        # Pattern matches "Model Name" followed by any whitespace, ":", more whitespace, and text
        model_pattern = re.search(r'(Model Name\s*:\s*)([^\n\r]+)', full_text, re.IGNORECASE)
        if model_pattern:
            old_value = model_pattern.group(2).strip()
            old_full = model_pattern.group(0)
            new_full = model_pattern.group(1) + new_model_name

            print(f"[PAC_GEN] Found Model Name: '{old_full}' in paragraph: {full_text[:100]}")
            if replace_in_paragraph(paragraph, old_full, new_full):
                print(f"[PAC_GEN] Replaced Model Name '{old_value}' with '{new_model_name}'")
                return True
        return False

    # Function to clear names but keep titles
    def clear_names_keep_titles(paragraph):
        """
        Clear names while keeping titles and formatting.
        Only clears signature table names, NOT customer name.
        """
        # Check if this paragraph contains "Name" field
        full_text = paragraph.text
        if 'Name' not in full_text or ':' not in full_text:
            return

        # Skip if this is "Customer Name" - we don't want to clear that
        if 'Customer' in full_text:
            return

        # Pattern to match "Name : <anything>" where <anything> is what we want to clear
        # This should only match signature fields like "Name :" not "Customer Name :"
        # Look for pattern that starts with just "Name" (possibly with "Signature" before it)
        name_pattern = r'^((?:Signature\s*:?\s*)?Name\s*:\s*)([^\n\r]+)'
        match = re.search(name_pattern, full_text, re.IGNORECASE)

        if not match:
            # Try alternative pattern for table cells with just "Name :"
            name_pattern = r'(^Name\s*:\s*)([^\n\r]+)'
            match = re.search(name_pattern, full_text, re.IGNORECASE)

        if not match:
            return

        # We found "Name : <value>", now we need to clear <value>
        # Build a map of text positions to runs
        name_value = match.group(2).strip()

        # Skip if the value is just underscores or dashes (already blank)
        if not name_value or re.match(r'^[_\-\s]+$', name_value):
            return

        print(f"[PAC_GEN] Found name to clear: '{name_value}' in: {full_text[:80]}")

        # Find the position where the name value starts
        value_start = match.start(2)
        value_end = match.end(2)

        # Map character positions to runs
        current_pos = 0
        for run in paragraph.runs:
            run_start = current_pos
            run_end = current_pos + len(run.text)

            # Check if this run contains part of the name value
            if run_start < value_end and run_end > value_start:
                # This run contains part of the name value
                overlap_start = max(0, value_start - run_start)
                overlap_end = min(len(run.text), value_end - run_start)

                # Clear the overlapping part
                if overlap_start == 0 and overlap_end == len(run.text):
                    # Clear the entire run
                    print(f"[PAC_GEN] Clearing entire run: '{run.text}'")
                    run.text = ""
                elif overlap_start == 0:
                    # Clear from start
                    run.text = run.text[overlap_end:]
                    print(f"[PAC_GEN] Cleared start of run, new: '{run.text}'")
                elif overlap_end == len(run.text):
                    # Clear to end
                    run.text = run.text[:overlap_start]
                    print(f"[PAC_GEN] Cleared end of run, new: '{run.text}'")
                else:
                    # Clear middle
                    run.text = run.text[:overlap_start] + run.text[overlap_end:]
                    print(f"[PAC_GEN] Cleared middle of run, new: '{run.text}'")

            current_pos = run_end

    # Print document structure info
    print(f"[PAC_GEN] Document has {len(doc.paragraphs)} paragraphs")
    print(f"[PAC_GEN] Document has {len(doc.tables)} tables")
    print(f"[PAC_GEN] Document has {len(doc.sections)} sections")

    # Check for images in the document
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
    print(f"[PAC_GEN] Document has {image_count} images")

    # Print first few paragraphs to see what's in the document
    for i, para in enumerate(doc.paragraphs[:20]):
        if para.text.strip():
            print(f"[PAC_GEN] Paragraph {i}: {para.text[:150]}")

    # Replace in document body
    print(f"[PAC_GEN] Processing document body paragraphs...")
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
        replace_po_line_number(paragraph, po_line_number)
        replace_model_name(paragraph, model_name)
        clear_names_keep_titles(paragraph)

    # Replace in tables
    print(f"[PAC_GEN] Processing tables (count: {len(doc.tables)})...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)
                    replace_po_line_number(paragraph, po_line_number)
                    replace_model_name(paragraph, model_name)
                    clear_names_keep_titles(paragraph)

    # Replace in headers and footers
    print(f"[PAC_GEN] Processing headers and footers...")
    for section in doc.sections:
        # Header
        header = section.header
        for paragraph in header.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
            replace_po_line_number(paragraph, po_line_number)
            replace_model_name(paragraph, model_name)
            clear_names_keep_titles(paragraph)

        # Footer
        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
            replace_po_line_number(paragraph, po_line_number)
            replace_model_name(paragraph, model_name)
            clear_names_keep_titles(paragraph)

    # Print summary of replacements
    print(f"[PAC_GEN] Replacement summary:")
    for old_text, count in replacements_made.items():
        print(f"[PAC_GEN]   '{old_text}': {count} replacements")

    # Save to BytesIO
    print(f"[PAC_GEN] Saving modified document...")
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    print(f"[PAC_GEN] Document saved successfully")

    # Optionally save to file
    if output_path:
        doc.save(output_path)

    return doc_bytes


def create_boq_zip_package(
    csv_content: str,
    site_id: str,
    project_name: str,
    project_po: str,
    link_id: str,
    template_path: str,
    csv_filename: str = "boq.csv",
    po_line_number: str = "1",
    model_name: str = "Implementation services - New Site"
) -> BytesIO:
    """
    Create a ZIP package containing BOQ CSV and modified PAC document.

    Args:
        csv_content: CSV content as string
        site_id: Site ID
        project_name: Project name
        project_po: Project PO number
        link_id: Link ID (for certificate number)
        template_path: Path to PAC template
        csv_filename: Name for the CSV file in the ZIP
        po_line_number: PO line number(s) - default "1"
        model_name: Model name from Implementation services row

    Returns:
        BytesIO object containing the ZIP file
    """
    print(f"[ZIP_PKG] Creating ZIP package...")
    print(f"[ZIP_PKG] CSV filename: {csv_filename}")
    print(f"[ZIP_PKG] CSV content length: {len(csv_content)}")
    print(f"[ZIP_PKG] Site ID: {site_id}")
    print(f"[ZIP_PKG] Link ID: {link_id}")
    print(f"[ZIP_PKG] Project name: {project_name}")
    print(f"[ZIP_PKG] Project PO: {project_po}")
    print(f"[ZIP_PKG] PO Line Number: {po_line_number}")
    print(f"[ZIP_PKG] Model Name: {model_name}")
    print(f"[ZIP_PKG] Template path: {template_path}")

    # Create ZIP file in memory
    zip_buffer = BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Add CSV file
        print(f"[ZIP_PKG] Adding CSV file to ZIP...")
        zip_file.writestr(csv_filename, csv_content)

        # Generate and add PAC document
        try:
            print(f"[ZIP_PKG] Generating PAC document...")
            pac_doc = modify_pac_template(
                template_path=template_path,
                site_id=site_id,
                project_name=project_name,
                project_po=project_po,
                link_id=link_id,
                po_line_number=po_line_number,
                model_name=model_name
            )

            # Add PAC document to ZIP
            pac_filename = f"PAC_{site_id}.docx"
            print(f"[ZIP_PKG] Adding PAC document to ZIP as: {pac_filename}")
            zip_file.writestr(pac_filename, pac_doc.getvalue())
            print(f"[ZIP_PKG] PAC document added successfully")
        except Exception as e:
            # If PAC generation fails, log error but still include CSV
            print(f"[ZIP_PKG] ERROR: Failed to generate PAC document: {str(e)}")
            import traceback
            traceback.print_exc()

    zip_buffer.seek(0)
    print(f"[ZIP_PKG] ZIP package created successfully")
    return zip_buffer
